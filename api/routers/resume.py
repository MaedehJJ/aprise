import io
import logging
import uuid
from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException, Request, status
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, ConfigDict
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from db.neon import get_db
from routers._limiter import limiter
from routers.auth import get_current_user
from services.ai_service import AIService, get_ai_service
from services.resume_service import ResumeService


class ATSScoreResponse(BaseModel):
    score: int
    matched_keywords: list[str]
    missing_keywords: list[str]
    formatting_issues: list[str]
    quick_fixes: list[str]

logger = logging.getLogger(__name__)
router = APIRouter()


class ExperienceEntryResponse(BaseModel):
    company: str
    role: str
    dates: str
    bullets: list[str]


class ResumeResponse(BaseModel):
    model_config = ConfigDict(from_attributes=True)

    id: uuid.UUID
    jd_id: uuid.UUID
    content: dict | None = None
    labels: dict | None = None
    is_generated: bool
    created_at: datetime


@router.post("/api/jds/{jd_id}/resume", response_model=ResumeResponse, status_code=201)
@limiter.limit("10/hour")
async def generate_resume(
    request: Request,
    jd_id: uuid.UUID,
    clerk_user_id: str = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    ai: AIService = Depends(get_ai_service),
):
    """
    Generates a tailored resume for the given JD.

    Pre-conditions:
      - A coaching conversation must exist for this JD.
      - The user must have memories ingested (CV upload).

    Post-effects:
      - Persists a new Resume row with structured content and retrieval tags.
      - Writes the generated tags back onto JD.labels for future similarity search.
    """
    service = ResumeService(db=db, ai=ai)
    resume = await service.generate_resume(jd_id=jd_id, clerk_user_id=clerk_user_id)
    return resume


@router.get("/api/jds/{jd_id}/resumes", response_model=list[ResumeResponse])
async def list_resumes(
    jd_id: uuid.UUID,
    clerk_user_id: str = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    ai: AIService = Depends(get_ai_service),
):
    service = ResumeService(db=db, ai=ai)
    return await service.list_resumes(jd_id=jd_id, clerk_user_id=clerk_user_id)


@router.get("/api/resumes/{resume_id}", response_model=ResumeResponse)
async def get_resume(
    resume_id: uuid.UUID,
    clerk_user_id: str = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    ai: AIService = Depends(get_ai_service),
):
    service = ResumeService(db=db, ai=ai)
    return await service.get_resume(resume_id=resume_id, clerk_user_id=clerk_user_id)


@router.get("/api/resumes/{resume_id}/pdf")
async def download_resume_pdf(
    resume_id: uuid.UUID,
    clerk_user_id: str = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    ai: AIService = Depends(get_ai_service),
):
    """
    Generate and return a formatted PDF for the given resume.
    The PDF is built from the structured JSONB content (summary, experience, skills).
    """
    from db.models import Resume as ResumeModel
    from services.utils import get_profile_or_404

    profile = await get_profile_or_404(db, clerk_user_id)
    resume = (
        await db.execute(
            select(ResumeModel)
            .options(selectinload(ResumeModel.jd))
            .filter_by(id=resume_id, user_id=profile.id)
        )
    ).scalars().unique().one_or_none()
    if not resume:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Resume not found.")

    if not resume.content:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Resume has no generated content yet.",
        )

    pdf_bytes = _build_resume_pdf(resume)
    filename = f"resume-{resume_id}.pdf"
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/api/resumes/{resume_id}/docx")
async def download_resume_docx(
    resume_id: uuid.UUID,
    clerk_user_id: str = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    ai: AIService = Depends(get_ai_service),
):
    """
    Return the stored DOCX for the given resume.
    If the DOCX has not been generated yet (legacy row), build it on the fly,
    persist it, and return it.
    """
    from db.models import Resume as ResumeModel
    from services.utils import get_profile_or_404

    profile = await get_profile_or_404(db, clerk_user_id)
    resume = (
        await db.execute(
            select(ResumeModel)
            .options(selectinload(ResumeModel.jd))
            .filter_by(id=resume_id, user_id=profile.id)
        )
    ).scalars().unique().one_or_none()
    if not resume:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Resume not found.")

    if not resume.content:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Resume has no generated content yet.",
        )

    if resume.docx_content:
        docx_bytes = resume.docx_content
    else:
        docx_bytes = _build_resume_docx(resume)
        resume.docx_content = docx_bytes
        await db.commit()

    filename = f"resume-{resume_id}.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/api/resumes/{resume_id}/ats-score", response_model=ATSScoreResponse)
@limiter.limit("20/hour")
async def get_ats_score(
    request: Request,
    resume_id: uuid.UUID,
    clerk_user_id: str = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    ai: AIService = Depends(get_ai_service),
):
    """
    Scores a resume's ATS compatibility against its source JD.
    Returns matched/missing keywords and quick fixes.
    """
    from db.models import Resume as ResumeModel
    from services.utils import get_profile_or_404
    from prompts import ats_score_prompt

    profile = await get_profile_or_404(db, clerk_user_id)
    resume = (
        await db.execute(
            select(ResumeModel)
            .options(selectinload(ResumeModel.jd))
            .filter_by(id=resume_id, user_id=profile.id)
        )
    ).scalars().unique().one_or_none()
    if not resume:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Resume not found.")
    if not resume.content:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Resume has no generated content.",
        )

    jd = resume.jd
    content = resume.content
    experience_text = "\n".join(
        f"{e.get('role', '')} at {e.get('company', '')} ({e.get('dates', '')}): "
        + " | ".join(e.get("bullets", [])[:3])
        for e in (content.get("experience") or [])
    )
    requirements = (jd.parsed_requirements or {}) if jd else {}

    result = ai.structured(
        ats_score_prompt,
        langsmith_extra={"resume_id": str(resume_id), "step": "ats_score"},
        company=jd.company_name if jd else "the company",
        role=jd.role_title if jd else "the role",
        required_skills=", ".join(requirements.get("required_skills", [])) or "Not specified",
        nice_to_have=", ".join(requirements.get("nice_to_have", [])) or "None",
        responsibilities="\n- " + "\n- ".join(requirements.get("responsibilities", [])) if requirements.get("responsibilities") else "Not specified",
        resume_summary=content.get("summary", ""),
        resume_experience=experience_text,
        resume_skills=", ".join(content.get("skills", [])),
    )
    return result


def _build_resume_docx(resume) -> bytes:
    """Build a clean, ATS-friendly DOCX from the structured resume content."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _set_color(run, hex_color: str):
        r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)

    def _add_horizontal_rule(doc):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "E5E7EB")
        pBdr.append(bottom)
        pPr.append(pBdr)

    content_data = resume.content or {}
    summary = content_data.get("summary", "")
    experience: list[dict] = content_data.get("experience", [])
    skills: list[str] = content_data.get("skills", [])

    jd = getattr(resume, "jd", None)
    role_title = (jd.role_title if jd else None) or "Tailored Resume"
    company = (jd.company_name if jd else None) or ""
    header_sub = f"Prepared for: {company} — {role_title}" if company else role_title

    doc = Document()

    # Narrow margins
    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    # Header
    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_after = Pt(2)
    run = title_para.add_run("Tailored Resume")
    run.bold = True
    run.font.size = Pt(18)
    _set_color(run, "1A1A2E")

    sub_para = doc.add_paragraph()
    sub_para.paragraph_format.space_after = Pt(8)
    run = sub_para.add_run(header_sub)
    run.font.size = Pt(10)
    _set_color(run, "6B7280")

    _add_horizontal_rule(doc)

    def _section_heading(text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        _set_color(run, "2563EB")

    # Summary
    if summary:
        _section_heading("PROFESSIONAL SUMMARY")
        p = doc.add_paragraph(summary)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(10)

    # Experience
    if experience:
        _section_heading("EXPERIENCE")
        _add_horizontal_rule(doc)
        for entry in experience:
            company_name = entry.get("company", "")
            role = entry.get("role", "")
            dates = entry.get("dates", "")
            bullets = entry.get("bullets", [])

            job_para = doc.add_paragraph()
            job_para.paragraph_format.space_before = Pt(6)
            job_para.paragraph_format.space_after = Pt(1)
            bold_run = job_para.add_run(f"{role} — {company_name}")
            bold_run.bold = True
            bold_run.font.size = Pt(10)
            date_run = job_para.add_run(f"  ({dates})")
            date_run.font.size = Pt(9)
            _set_color(date_run, "6B7280")

            for bullet in bullets:
                b_para = doc.add_paragraph(style="List Bullet")
                b_para.paragraph_format.space_after = Pt(1)
                b_para.paragraph_format.left_indent = Inches(0.2)
                run = b_para.add_run(bullet)
                run.font.size = Pt(9.5)

    # Skills
    if skills:
        _section_heading("SKILLS")
        _add_horizontal_rule(doc)
        p = doc.add_paragraph(" • ".join(skills))
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_resume_pdf(resume) -> bytes:
    """Build a clean, single-page PDF from the structured resume content."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        HRFlowable, ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()

    # ── Custom styles ──────────────────────────────────────────────────
    heading1 = ParagraphStyle(
        "Heading1",
        parent=styles["Heading1"],
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#1a1a2e"),
        spaceAfter=4,
    )
    section_title = ParagraphStyle(
        "SectionTitle",
        parent=styles["Normal"],
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#2563eb"),
        fontName="Helvetica-Bold",
        spaceBefore=12,
        spaceAfter=4,
    )
    job_title_style = ParagraphStyle(
        "JobTitle",
        parent=styles["Normal"],
        fontSize=10,
        leading=13,
        fontName="Helvetica-Bold",
        textColor=colors.HexColor("#111827"),
        spaceBefore=6,
    )
    body = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor("#374151"),
    )
    bullet_style = ParagraphStyle(
        "Bullet",
        parent=body,
        leftIndent=14,
        bulletIndent=4,
        spaceAfter=1,
    )

    content_data = resume.content or {}
    summary = content_data.get("summary", "")
    experience: list[dict] = content_data.get("experience", [])
    skills: list[str] = content_data.get("skills", [])

    jd = getattr(resume, "jd", None)
    role_title = (jd.role_title if jd else None) or "Tailored Resume"
    company = (jd.company_name if jd else None) or ""
    header_sub = f"Prepared for: {company} — {role_title}" if company else role_title

    story = []

    # ── Header ────────────────────────────────────────────────────────
    story.append(Paragraph("Tailored Resume", heading1))
    story.append(Paragraph(header_sub, ParagraphStyle(
        "SubHead", parent=body, textColor=colors.HexColor("#6b7280"), spaceAfter=8
    )))
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#2563eb")))

    # ── Summary ───────────────────────────────────────────────────────
    if summary:
        story.append(Paragraph("PROFESSIONAL SUMMARY", section_title))
        story.append(Paragraph(summary, body))

    # ── Experience ────────────────────────────────────────────────────
    if experience:
        story.append(Paragraph("EXPERIENCE", section_title))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#e5e7eb"), spaceAfter=4))
        for entry in experience:
            co = entry.get("company", "")
            ro = entry.get("role", "")
            dt = entry.get("dates", "")
            bullets = entry.get("bullets", [])
            story.append(Paragraph(
                f"<b>{ro}</b> — {co} <font color='#6b7280' size='9'>({dt})</font>",
                job_title_style,
            ))
            for b in bullets:
                story.append(Paragraph(f"• {b}", bullet_style))
            story.append(Spacer(1, 4))

    # ── Skills ────────────────────────────────────────────────────────
    if skills:
        story.append(Paragraph("SKILLS", section_title))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#e5e7eb"), spaceAfter=4))
        story.append(Paragraph(" • ".join(skills), body))

    doc.build(story)
    return buf.getvalue()
